# -*- coding: utf-8 -*-
"""
Archivo: wordgen.py
Descripción: Generación del documento Word para incidencias. El script confía en
             que el input del lugar ya incluye el artículo (ej. "el patio").
"""

import os
import random
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

from config import SCHOOL_NAME, LOCATION, DIRECTOR_NAME, TEACHER_NAME, GRADE, GROUP

# El diccionario vuelve a usar la variable {lugar} directamente.
SINONIMOS = {
    'apertura': [
        "Siendo las {hora} horas del día {fecha},",
        "En la fecha {fecha}, alrededor de las {hora} horas,",
        "Aproximadamente a las {hora} horas del día {fecha},"
    ],
    'contexto': [
        "durante {actividad} en {lugar},",
        "mientras se realizaba la actividad '{actividad}' en {lugar},",
        "en el transcurso de {actividad}, llevada a cabo en {lugar},"
    ],
    'suceso': [
        "se presentó una incidencia del tipo {tipo_inc}.",
        "ocurrió un suceso clasificado como {tipo_inc}.",
        "se registró un incidente de tipo {tipo_inc}.",
        "tuvo lugar un evento del tipo {tipo_inc}."
    ],
    'participantes': [
        # Se eliminó la palabra "grupo"; queda "del {GRADE}"{GROUP}"" tal como pediste.
        'En ella participaron {participantes_str} del {GRADE}"{GROUP}".',
        'En el suceso estuvieron involucrados {participantes_str} del {GRADE}"{GROUP}".',
        'Los participantes fueron {participantes_str} del {GRADE}"{GROUP}".'
    ],
    'gravedad': [
        "La gravedad fue evaluada como {gravedad_lower}.",
        "El incidente fue catalogado con una gravedad {gravedad_lower}.",
        "Se determinó que la seriedad del evento fue {gravedad_lower}."
    ],
    'descripcion_hechos': [
        "Los hechos se describen de la siguiente manera: {narracion}.",
        "A continuación se narra lo sucedido: {narracion}.",
        "La descripción de los acontecimientos es la siguiente: {narracion}."
    ]
}


def set_cell_borders(cell, **kwargs):
    """
    Función para establecer bordes específicos en una celda.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = OxmlElement(f'w:{edge}')
            for key, val in edge_data.items():
                tag.set(qn(f'w:{key}'), str(val))
            tcBorders.append(tag)

    tcPr.append(tcBorders)


def generar_word(fecha, hora, lugar, actividad, participantes, tipo_inc,
                 gravedad, narracion, medidas, seguimiento, padres_dict,
                 alumnos_seleccionados, output_path):
    """
    Genera el documento Word de la bitácora de manera segura.
    """
    if not isinstance(padres_dict, dict):
        padres_dict = {}

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)

    margin_size = Inches(0.5)
    sec.top_margin = margin_size
    sec.bottom_margin = margin_size
    sec.left_margin = margin_size
    sec.right_margin = margin_size

    # ----- Encabezado con logos -----
    try:
        header = sec.header
        header_table = header.add_table(rows=1, cols=3, width=Inches(7.5))
        header_table.autofit = False

        cell_logo1 = header_table.cell(0, 0)
        cell_logo1.width = Inches(1.5)
        cell_logo1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo1_path = os.path.join("recursos", "logo1.png")
        if os.path.exists(logo1_path):
            run = cell_logo1.paragraphs[0].add_run()
            run.add_picture(logo1_path, width=Inches(1.0))
        else:
            cell_logo1.paragraphs[0].add_run("Logo Izquierdo")

        cell_center = header_table.cell(0, 1)
        cell_center.width = Inches(3.5)

        cell_logo2 = header_table.cell(0, 2)
        cell_logo2.width = Inches(2.5)
        cell_logo2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo2_path = os.path.join("recursos", "logo2.png")
        if os.path.exists(logo2_path):
            run = cell_logo2.paragraphs[0].add_run()
            run.add_picture(logo2_path, width=Inches(2.5))
        else:
            cell_logo2.paragraphs[0].add_run("Logo Derecho")
    except Exception as e:
        print(f"Advertencia: No se pudieron agregar los logos al encabezado. Causa: {e}")

    # ----- Título -----
    t = doc.add_paragraph()
    r = t.add_run(f"BITÁCORA DE INCIDENCIA - {SCHOOL_NAME}\n")
    r.bold = True
    r.font.size = Pt(14)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p_sub = doc.add_paragraph(LOCATION, style='Subtitle')
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        # Si style 'Subtitle' no existe, añadimos sin estilo
        p = doc.add_paragraph(LOCATION)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ----- Narración Dinámica con formato -----
    participantes_str = ', '.join(participantes)

    frase_apertura = random.choice(SINONIMOS['apertura']).format(hora=hora, fecha=fecha)
    frase_contexto = random.choice(SINONIMOS['contexto']).format(actividad=actividad, lugar=lugar)
    frase_suceso = random.choice(SINONIMOS['suceso']).format(tipo_inc=tipo_inc)
    frase_participantes = random.choice(SINONIMOS['participantes']).format(participantes_str=participantes_str, GRADE=GRADE, GROUP=GROUP)
    frase_gravedad = random.choice(SINONIMOS['gravedad']).format(gravedad_lower=gravedad.lower())
    frase_descripcion = random.choice(SINONIMOS['descripcion_hechos']).format(narracion=narracion)

    narr = (
        f"{frase_apertura} {frase_contexto} {frase_suceso} "
        f"{frase_participantes} {frase_gravedad} {frase_descripcion}"
    )

    if medidas:
        narr += f" Las medidas tomadas fueron: {medidas}."
    if seguimiento:
        narr += f" Para su seguimiento se determinó: {seguimiento}."

    p_narr = doc.add_paragraph()
    run_narr = p_narr.add_run(narr)
    run_narr.font.size = Pt(12)
    p_narr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # ----- Tabla de Firmas Estilizada -----
    firmas_data = []
    if gravedad in ["Moderada", "Grave"]:
        firmas_data.append(("Director", DIRECTOR_NAME))
    firmas_data.append(("Maestro de Grupo", TEACHER_NAME))
    for alumno in alumnos_seleccionados:
        firmas_data.append(("Alumno", alumno))
    if gravedad == "Grave":
        for alumno in alumnos_seleccionados:
            padre = str(padres_dict.get(alumno, "Padre/Madre de familia"))
            firmas_data.append((f"Padre/Madre de familia ({alumno})", padre))
    firmas_data.append(("Testigo", ""))

    signatures_table = doc.add_table(rows=1, cols=3)
    signatures_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = (Inches(2.0), Inches(3.0), Inches(2.5))

    border_normal = {"sz": 6, "val": "single", "color": "000000"}
    border_thick = {"sz": 12, "val": "single", "color": "000000"}
    border_none = {"val": "nil"}

    hdr_cells = signatures_table.rows[0].cells
    hdr_cells[0].text = 'Cargo / Relación'
    hdr_cells[1].text = 'Nombre Completo'
    hdr_cells[2].text = 'Firma'

    for i, cell in enumerate(hdr_cells):
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.width = col_widths[i]
        set_cell_borders(cell, top=border_normal, bottom=border_thick, left=border_none, right=border_none)

    for cargo, nombre in firmas_data:
        row_cells = signatures_table.add_row().cells
        row_cells[0].text = cargo
        row_cells[1].text = nombre
        row_cells[2].text = ''

        for i, cell in enumerate(row_cells):
            cell.width = col_widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell, bottom=border_normal, top=border_none, left=border_none, right=border_none)

    # Guardar documento
    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    return output_path
